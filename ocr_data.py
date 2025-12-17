import os
import re
import torch
import numpy as np
import gc
import argparse
from typing import List, Dict
from PIL import Image
from pdf2image import convert_from_path, pdfinfo_from_path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm

# --- IMPORT MODELS ---
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg
from easyocr import Reader

# --- CONFIGURATION ---
Image.MAX_IMAGE_PIXELS = None  # Prevent DecompressionBombError for large images

class Config:
    """
    Global configuration settings for the OCR pipeline.
    """
    DPI_RESOLUTION = 300  # High DPI for better OCR accuracy
    TOLERANCE_Y = 30      # Vertical tolerance (pixels) for line merging
    MARGIN_HEADER = 0.05  # Ignore top 5% (Header)
    MARGIN_FOOTER = 0.05   # Ignore bottom 5% (Footer)
    BATCH_SIZE = 10       # Process 10 pages at a time to save RAM
    
    # [UPDATED REGEX]
    # 1. (\d+\.\d+) : Mandatory Number.Number format (Level 2). E.g., 2.1
    # 2. \s* : Accept extra whitespace (OCR artifacts).
    # 3. \.?        : Dot is OPTIONAL (Accepts "2.1." or "2.1").
    # 4. \s+        : MANDATORY space afterwards (prevents matching 2.1.1).
    REGEX_HEADER_L2 = re.compile(r"^\s*(\d+\.\d+)\s*\.?\s+(?P<content>.*)$")


# --- INIT DEVICE ---
COMPUTE_DEVICE = "cuda:0" if torch.cuda.is_available() else "cpu"
HAS_CUDA = torch.cuda.is_available()

print(f"System initialized on device: {COMPUTE_DEVICE.upper()}")

# --- LOAD MODELS ---
print("Initializing VietOCR engine...")
voc_cfg = Cfg.load_config_from_name("vgg_transformer")
voc_cfg["cnn"]["pretrained"] = False
voc_cfg["device"] = COMPUTE_DEVICE
voc_cfg["predictor"]["beamsearch"] = False
text_recognizer = Predictor(voc_cfg)

print("Initializing EasyOCR engine...")
text_detector = Reader(["vi"], gpu=HAS_CUDA, quantize=False)


def add_padding(img_np: np.ndarray, pad: int = 10) -> np.ndarray:
    """
    Adds white padding around an image crop.
    
    Why: VietOCR performs significantly better when text is not touching the edges.
    
    Args:
        img_np (np.ndarray): Input image array (H, W, C).
        pad (int): Number of pixels to pad on all sides.
        
    Returns:
        np.ndarray: Padded image array.
    """
    h, w = img_np.shape[:2]
    padded = np.ones((h + 2*pad, w + 2*pad, 3), dtype=np.uint8) * 255
    padded[pad:h+pad, pad:w+pad] = img_np
    return padded


def validate_header(text):
    """
    Validates if a line is a Level 2 Header (e.g., 2.1 or 2.1.).
    Excludes Level 1 (2.) and Level 3 (2.1.1).
    """
    if len(text) < 4: return False
    clean = text.strip()
    
    # 1. Blacklist check (Skip Figure/Table captions)
    # "Hình", "Bảng", "Sơ đồ" -> Figure, Table, Diagram
    if any(clean.lower().startswith(x) for x in ["Hình", "Hinh", "Bảng", "Bang", "Sơ đồ", "Biểu đồ"]):
        return False

    # 2. Check Regex Pattern
    match = Config.REGEX_HEADER_L2.match(clean)
    if not match: 
        return False 
    
    # 3. Double Check (Count numbers to ensure Level 3 is excluded)
    # Get the captured number part (E.g., "2.1")
    number_part = match.group(1)
    # Check if there are exactly 2 numbers
    if len(number_part.split('.')) != 2:
        return False

    # 4. Filter measurement units (Avoid mistaking "1.5 mg" for a header)
    content = match.group("content")
    if re.search(r"^(mg|g|ml|lít|cm|mm|%|độ|kg)(\s|$|\W)", content, re.IGNORECASE): 
        return False
        
    # 5. Ensure there is alphabetic content in the header title
    if not re.search(r"[a-zA-Z]", content): 
        return False
    
    return True


# ==============================================================================
# 2. OCR ENGINE 
# ==============================================================================

def ocr_process_page(img_array: np.ndarray) -> List[Dict]:
    """
    Performs OCR on a single page image.
    
    Pipeline:
        1. Detect text bounding boxes using EasyOCR.
        2. Filter out boxes in Header/Footer regions.
        3. Crop and Recognize text using VietOCR.
        4. Sort and Merge disjointed words into full lines.
        
    Args:
        img_array (np.ndarray): The full page image.
        
    Returns:
        List[Dict]: A list of text blocks with content and coordinates.
    """
    img_h, img_w = img_array.shape[:2]
    
    # Step 1: Detect text boxes
    try:
        results = text_detector.readtext(img_array, detail=1, width_ths=0.7)
    except Exception as e:
        print(f"[WARN] Detection failed: {e}")
        return []

    raw_blocks = []
    y_min_limit = img_h * Config.MARGIN_HEADER
    y_max_limit = img_h * (1 - Config.MARGIN_FOOTER)

    # Step 2: Filter and Recognize
    for box, _, conf in results:
        # Confidence Threshold
        if conf < 0.25: continue
        
        y_coords = [p[1] for p in box]
        y_center = sum(y_coords) / 4
        
        # Filter Header/Footer
        if y_center < y_min_limit or y_center > y_max_limit: continue

        xs = [p[0] for p in box]
        ys = [p[1] for p in box]
        x1, x2 = int(min(xs)), int(max(xs))
        y1, y2 = int(min(ys)), int(max(ys))

        # Expand Crop slightly (Fix lost diacritics like 'Ở')
        crop_expand = 5
        crop = img_array[
            max(0, y1-crop_expand):min(img_h, y2+crop_expand), 
            max(0, x1-crop_expand):min(img_w, x2+crop_expand)
        ]
        
        if crop.size == 0: continue

        try:
            # Padding + Recognize
            padded_crop = add_padding(crop, pad=15)
            text = text_recognizer.predict(Image.fromarray(padded_crop))
            
            if len(text.strip()) > 1:
                raw_blocks.append({
                    "text": text.strip(),
                    "y_center": y_center / img_h, # Normalized Y for sorting
                    "x_min": x1
                })
        except: continue

    if not raw_blocks: return []

    # Step 3: Line Merging Algorithm (Stable Logic)
    # Sort vertically first
    raw_blocks.sort(key=lambda item: item["y_center"])
    
    lines = []
    current_line = [raw_blocks[0]]
    
    # Use normalized tolerance ratio
    tolerance_ratio = Config.TOLERANCE_Y / img_h

    for i in range(1, len(raw_blocks)):
        current = raw_blocks[i]
        prev = current_line[-1]
        
        # Check vertical distance to merge into same line
        if abs(current["y_center"] - prev["y_center"]) <= tolerance_ratio:
            current_line.append(current)
        else:
            lines.append(current_line)
            current_line = [current]
    lines.append(current_line)

    # Step 4: Construct Final Output
    final_output = []
    for line in lines:
        # Sort words left-to-right within the line
        line.sort(key=lambda k: k["x_min"])
        joined_text = " ".join([b["text"] for b in line])
        
        final_output.append({
            "content": joined_text
        })
        
    return final_output


# ==============================================================================
# 3. EXPORT TO WORD
# ==============================================================================

def save_to_word(data_list: List[Dict], output_path: str):
    """
    Saves the extracted text to a .docx file.
    Inserts </break> tags before Level 2 Headers.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    
    buffer = []
    
    def flush(buf):
        for item in buf:
            doc.add_paragraph(item["content"])
            
    for item in data_list:
        text = item["content"]
        
        # Identify Header to Insert Break
        if validate_header(text):
            print(f"[INFO] New Section Detected: {text[:40]}...")
            
            if buffer:
                flush(buffer)
                
                # Insert colored Break Tag
                p = doc.add_paragraph()
                run = p.add_run("</break>")
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                buffer = []
            buffer.append(item)
        else:
            buffer.append(item)
            
    # Flush remaining content
    if buffer: flush(buffer)
    doc.save(output_path)


# ==============================================================================
# 4. MAIN PIPELINE
# ==============================================================================

def run_pipeline(pdf_path: str, start: int = 1, end: int = None):
    """
    Orchestrates the entire OCR process for a single PDF file.
    
    Args:
        pdf_path (str): Path to input PDF.
        start (int): Start page number.
        end (int): End page number (None for all).
    """
    if not os.path.exists(pdf_path):
        print(f"[ERROR] File not found: {pdf_path}")
        return

    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    out_dir = os.path.join("output", f"{base_name}")
    os.makedirs(out_dir, exist_ok=True)
    
    print(f"[INFO] Processing Target: {base_name}")
    
    # Get page count
    try:
        info = pdfinfo_from_path(pdf_path)
        total = info["Pages"]
        print(f"[INFO] Total pages: {total}")
    except: 
        total = 1000
        print("[WARN] Could not read PDF metadata. Defaulting to safe limit.")
    
    target_end = min(end, total) if end else total
    all_data = []
    
    # Batch Processing Loop
    for b_start in range(start, target_end + 1, Config.BATCH_SIZE):
        b_end = min(b_start + Config.BATCH_SIZE - 1, target_end)
        print(f"\n[INFO] Processing Batch: {b_start} -> {b_end}")
        
        try:
            images = convert_from_path(pdf_path, dpi=Config.DPI_RESOLUTION, first_page=b_start, last_page=b_end)
        except Exception as e:
            print(f"[ERROR] PDF Conversion failed: {e}")
            continue
            
        for img in images:
            # TEXT ONLY processing
            page_text = ocr_process_page(np.asarray(img))
            all_data.extend(page_text)
            
        del images
        gc.collect()
        
    out_name = f"{base_name}.docx"
    save_path = os.path.join(out_dir, out_name)
    save_to_word(all_data, save_path)
    
    print("-" * 50)
    print(f"[SUCCESS] Workflow Completed.")
    print(f"[OUTPUT] Saved to: {save_path}")
    print("-" * 50)


# ==============================================================================
# ENTRY POINT
# ==============================================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Medical OCR Tool")
    parser.add_argument("--input", type=str, help="Path to specific PDF file")
    parser.add_argument("--start", type=int, default=1, help="Start page number")
    parser.add_argument("--end", type=int, help="End page number")
    
    args = parser.parse_args()
    
    if args.input:
        run_pipeline(args.input, args.start, args.end)
    else:
        # Auto-scan 'input' directory
        target = "input"
        if os.path.exists(target):
            files = [f for f in os.listdir(target) if f.lower().endswith(".pdf")]
            if not files:
                print(f"[INFO] No PDF files found in '{target}'")
            else:
                print(f"[INFO] Found {len(files)} files. Starting batch processing...")
                for f in files:
                    run_pipeline(os.path.join(target, f), args.start, args.end)
        else:
            print(f"[ERROR] Directory '{target}' not found. Please create it or use --input.")


# python ocr_data.py --input "input/dieu-tri-hoc-ket-hop-y-hoc-hien-dai-va-y-hoc-co-truyen.pdf"